import boto3
import logging
from typing import List, Dict, Any
from botocore.exceptions import ClientError, BotoCoreError
from config import settings
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
import tempfile
import os
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from io import BytesIO

# Configure logging
logging.basicConfig(
    level=getattr(logging, settings.log_level),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class CloudFrontCrossAccountManager:
    """Manages cross-account CloudFront distribution listing using AWS STS AssumeRole"""
    
    def __init__(self):
        self.source_account_id = settings.source_account_id
        self.target_account_id = settings.target_account_id
        self.assume_role_name = settings.assume_role_name
        self.session_name = settings.session_name
        self.region = settings.aws_region
        
    def assume_role(self) -> Dict[str, Any]:
        """
        Assume role in the target account
        
        Returns:
            Dict containing temporary credentials
        """
        try:
            # Create STS client using instance profile credentials
            sts_client = boto3.client('sts', region_name=self.region)
            
            # Construct the role ARN
            role_arn = f"arn:aws:iam::{self.target_account_id}:role/{self.assume_role_name}"
            
            logger.info(f"Attempting to assume role: {role_arn}")
            
            # Assume the role
            response = sts_client.assume_role(
                RoleArn=role_arn,
                RoleSessionName=self.session_name,
                DurationSeconds=3600,  # 1 hour
                ExternalId='starbucks-monitoring-secret-key-prod'  # Uncomment if using ExternalId
            )
            
            logger.info("Successfully assumed role in target account")
            return response['Credentials']
            
        except ClientError as e:
            error_code = e.response['Error']['Code']
            error_message = e.response['Error']['Message']
            logger.error(f"Failed to assume role: {error_code} - {error_message}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error during assume role: {str(e)}")
            raise
    
    def create_cloudfront_client(self, credentials: Dict[str, Any]):
        """
        Create CloudFront client with assumed role credentials
        
        Args:
            credentials: Temporary credentials from assume_role
            
        Returns:
            CloudFront client object
        """
        # CloudFront is a global service, so we use us-east-1
        return boto3.client(
            'cloudfront',
            region_name='us-east-1',
            aws_access_key_id=credentials['AccessKeyId'],
            aws_secret_access_key=credentials['SecretAccessKey'],
            aws_session_token=credentials['SessionToken']
        )
    
    def create_cloudwatch_client(self, credentials: Dict[str, Any]):
        """
        Create CloudWatch client with assumed role credentials
        
        Args:
            credentials: Temporary credentials from assume_role
            
        Returns:
            CloudWatch client object
        """
        # CloudWatch for CloudFront metrics is in us-east-1
        return boto3.client(
            'cloudwatch',
            region_name='us-east-1',
            aws_access_key_id=credentials['AccessKeyId'],
            aws_secret_access_key=credentials['SecretAccessKey'],
            aws_session_token=credentials['SessionToken']
        )
    
    def list_cloudfront_distributions(self) -> List[Dict[str, Any]]:
        """
        List all CloudFront distributions in the target account
        
        Returns:
            List of CloudFront distribution information
        """
        try:
            # Assume role in target account
            credentials = self.assume_role()
            
            # Create CloudFront client with assumed credentials
            cf_client = self.create_cloudfront_client(credentials)
            
            # List all distributions
            logger.info(f"Retrieving CloudFront distributions from account {self.target_account_id}")
            
            distributions = []
            paginator = cf_client.get_paginator('list_distributions')
            
            for page in paginator.paginate():
                if 'Items' in page['DistributionList']:
                    for distribution in page['DistributionList']['Items']:
                        # Get detailed distribution information
                        dist_id = distribution['Id']
                        
                        # Get tags for the distribution
                        try:
                            tags_response = cf_client.list_tags_for_resource(Resource=f"arn:aws:cloudfront::{self.target_account_id}:distribution/{dist_id}")
                            tags = self._parse_tags(tags_response.get('Tags', {}).get('Items', []))
                        except Exception as tag_error:
                            logger.warning(f"Failed to get tags for distribution {dist_id}: {tag_error}")
                            tags = {}
                        
                        # Extract origin information
                        origins = []
                        for origin in distribution.get('Origins', {}).get('Items', []):
                            origin_info = {
                                'Id': origin.get('Id'),
                                'DomainName': origin.get('DomainName'),
                                'OriginPath': origin.get('OriginPath', ''),
                            }
                            # Check origin type
                            if 'S3OriginConfig' in origin:
                                origin_info['Type'] = 'S3'
                                origin_info['OAI'] = origin['S3OriginConfig'].get('OriginAccessIdentity', '')
                            elif 'CustomOriginConfig' in origin:
                                origin_info['Type'] = 'Custom'
                                origin_info['Protocol'] = origin['CustomOriginConfig'].get('OriginProtocolPolicy', '')
                            
                            origins.append(origin_info)
                        
                        distribution_info = {
                            'Id': dist_id,
                            'ARN': distribution.get('ARN'),
                            'DomainName': distribution.get('DomainName'),
                            'Status': distribution.get('Status'),
                            'Enabled': distribution.get('Enabled'),
                            'Comment': distribution.get('Comment', ''),
                            'PriceClass': distribution.get('PriceClass'),
                            'HttpVersion': distribution.get('HttpVersion'),
                            'IsIPV6Enabled': distribution.get('IsIPV6Enabled'),
                            'WebACLId': distribution.get('WebACLId', ''),
                            'LastModifiedTime': distribution.get('LastModifiedTime').isoformat() if distribution.get('LastModifiedTime') else '',
                            'Origins': origins,
                            'DefaultRootObject': distribution.get('DefaultRootObject', ''),
                            'Aliases': distribution.get('Aliases', {}).get('Items', []),
                            'Tags': tags
                        }
                        distributions.append(distribution_info)
            
            logger.info(f"Successfully retrieved {len(distributions)} CloudFront distributions")
            return distributions
            
        except Exception as e:
            logger.error(f"Failed to list CloudFront distributions: {str(e)}")
            raise
    
    def list_cloudfront_origin_access_identities(self) -> List[Dict[str, Any]]:
        """
        List all CloudFront Origin Access Identities in the target account
        
        Returns:
            List of OAI information
        """
        try:
            # Assume role in target account
            credentials = self.assume_role()
            
            # Create CloudFront client with assumed credentials
            cf_client = self.create_cloudfront_client(credentials)
            
            # List all OAIs
            logger.info(f"Retrieving CloudFront OAIs from account {self.target_account_id}")
            
            oais = []
            paginator = cf_client.get_paginator('list_cloud_front_origin_access_identities')
            
            for page in paginator.paginate():
                if 'Items' in page['CloudFrontOriginAccessIdentityList']:
                    for oai in page['CloudFrontOriginAccessIdentityList']['Items']:
                        oai_info = {
                            'Id': oai.get('Id'),
                            'S3CanonicalUserId': oai.get('S3CanonicalUserId'),
                            'Comment': oai.get('Comment', '')
                        }
                        oais.append(oai_info)
            
            logger.info(f"Successfully retrieved {len(oais)} CloudFront OAIs")
            return oais
            
        except Exception as e:
            logger.error(f"Failed to list CloudFront OAIs: {str(e)}")
            raise
    
    def get_cloudfront_metrics(self, distribution_id: str, metric_name: str, start_time: datetime, end_time: datetime) -> Dict[str, Any]:
        """
        Get CloudFront distribution metrics from CloudWatch
        
        Args:
            distribution_id: CloudFront distribution ID
            metric_name: Name of the metric (Requests or BytesDownloaded)
            start_time: Start time for metrics
            end_time: End time for metrics
            
        Returns:
            Dict containing datapoints, max_value, and max_timestamp
        """
        try:
            # Assume role in target account
            credentials = self.assume_role()
            
            # Create CloudWatch client with assumed credentials
            cw_client = self.create_cloudwatch_client(credentials)
            
            logger.info(f"Retrieving {metric_name} metrics for distribution {distribution_id}")
            
            # Get metric statistics
            response = cw_client.get_metric_statistics(
                Namespace='AWS/CloudFront',
                MetricName=metric_name,
                Dimensions=[
                    {
                        'Name': 'DistributionId',
                        'Value': distribution_id
                    },
                    {
                        'Name': 'Region',
                        'Value': 'Global'
                    }
                ],
                StartTime=start_time,
                EndTime=end_time,
                Period=3600,  # 1 hour
                Statistics=['Sum']
            )
            
            # Sort data points by timestamp
            datapoints = sorted(response['Datapoints'], key=lambda x: x['Timestamp'])
            
            # Find max value and timestamp
            max_value = None
            max_timestamp = None
            
            if datapoints:
                max_point = max(datapoints, key=lambda x: x['Sum'])
                max_value = max_point['Sum']
                max_timestamp = max_point['Timestamp']
            
            result = {
                'datapoints': datapoints,
                'max_value': max_value,
                'max_timestamp': max_timestamp
            }
            
            logger.info(f"Retrieved {len(datapoints)} data points for {metric_name}")
            if max_value is not None:
                logger.info(f"Max value: {max_value} at {max_timestamp}")
            
            return result
            
        except Exception as e:
            logger.error(f"Failed to get CloudFront metrics: {str(e)}")
            return {'datapoints': [], 'max_value': None, 'max_timestamp': None}
    
    def generate_metric_graph(self, distribution_id: str, metric_name: str, datapoints: List[Dict[str, Any]], max_value: float = None, max_timestamp: datetime = None, unit: str = '') -> BytesIO:
        """
        Generate a graph for CloudFront metrics
        
        Args:
            distribution_id: CloudFront distribution ID
            metric_name: Name of the metric
            datapoints: List of metric data points
            max_value: Maximum value in the dataset
            max_timestamp: Timestamp when maximum value occurred
            unit: Unit for the metric (e.g., 'Bytes', 'Count')
            
        Returns:
            BytesIO object containing the graph image
        """
        try:
            # Create figure and axis
            fig, ax = plt.subplots(figsize=(10, 6))
            
            if datapoints:
                # Extract timestamps and values
                timestamps = [point['Timestamp'] for point in datapoints]
                values = [point['Sum'] for point in datapoints]
                
                # Plot the data
                ax.plot(timestamps, values, marker='o', linestyle='-', linewidth=2, markersize=4)
                
                # Highlight max value point if available
                if max_value is not None and max_timestamp is not None:
                    ax.plot(max_timestamp, max_value, marker='*', markersize=12, color='red', 
                           label=f'Max: {self._format_metric_value(max_value, metric_name)} at {max_timestamp.strftime("%m-%d %H:%M")}')
                    
                    # Add annotation for max point
                    ax.annotate(f'Max: {self._format_metric_value(max_value, metric_name)}', 
                               xy=(max_timestamp, max_value), 
                               xytext=(10, 10), 
                               textcoords='offset points',
                               bbox=dict(boxstyle='round,pad=0.3', facecolor='yellow', alpha=0.7),
                               arrowprops=dict(arrowstyle='->', connectionstyle='arc3,rad=0'))
                
                # Format the x-axis
                ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d %H:%M'))
                ax.xaxis.set_major_locator(mdates.HourLocator(interval=24))
                plt.xticks(rotation=45, ha='right')
                
                # Add grid
                ax.grid(True, alpha=0.3)
                
                # Set labels and title
                ax.set_xlabel('Date/Time', fontsize=10)
                ylabel = f'{metric_name}'
                if unit:
                    ylabel += f' ({unit})'
                ax.set_ylabel(ylabel, fontsize=10)
                ax.set_title(f' ', fontsize=12, fontweight='bold')
                
                # Format y-axis values
                if metric_name == 'BytesDownloaded':
                    # Convert to human-readable format (GB)
                    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'{x/1e9:.2f} GB'))
                elif metric_name == 'Requests':
                    # Format as thousands
                    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'{x/1000:.0f}K' if x >= 1000 else f'{x:.0f}'))
                
                # Add legend if max value is shown
                if max_value is not None:
                    ax.legend(loc='upper right', fontsize=8)
                    
            else:
                # No data available
                ax.text(0.5, 0.5, 'No data available', 
                       horizontalalignment='center',
                       verticalalignment='center',
                       transform=ax.transAxes,
                       fontsize=14)
                ax.set_title(f' ', fontsize=12, fontweight='bold')
            
            # Adjust layout
            plt.tight_layout()
            
            # Save to BytesIO
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', dpi=100, bbox_inches='tight')
            img_buffer.seek(0)
            
            # Clean up
            plt.close(fig)
            
            return img_buffer
            
        except Exception as e:
            logger.error(f"Failed to generate metric graph: {str(e)}")
            # Return empty buffer on error
            return BytesIO()
    
    def _format_metric_value(self, value: float, metric_name: str) -> str:
        """
        Format metric value for display
        
        Args:
            value: The metric value
            metric_name: Name of the metric
            
        Returns:
            Formatted string representation
        """
        if metric_name == 'BytesDownloaded':
            return f'{value/1e9:.2f} GB'
        elif metric_name == 'Requests':
            if value >= 1000:
                return f'{value/1000:.0f}K'
            else:
                return f'{value:.0f}'
        else:
            return str(value)
    
    def _parse_tags(self, tags: List[Dict[str, str]]) -> Dict[str, str]:
        """
        Parse CloudFront tags into a dictionary
        
        Args:
            tags: List of tag dictionaries from CloudFront
            
        Returns:
            Dictionary of tag key-value pairs
        """
        return {tag.get('Key', ''): tag.get('Value', '') for tag in tags}
    
    def _safe_str(self, value: Any) -> str:
        """
        Safely convert any value to string, handling None values
        
        Args:
            value: Any value that needs to be converted to string
            
        Returns:
            String representation of the value, or empty string if None
        """
        if value is None:
            return ''
        if isinstance(value, bool):
            return 'Yes' if value else 'No'
        if isinstance(value, list):
            return ', '.join(str(item) for item in value)
        return str(value)
    
    def generate_cloudfront_word_report(self, distributions: List[Dict[str, Any]], oais: List[Dict[str, Any]]) -> str:
        """
        Generate Word document report with CloudFront data and CloudWatch metrics graphs
        
        Args:
            distributions: List of CloudFront distribution information
            oais: List of Origin Access Identity information
            
        Returns:
            Path to generated Word document
        """
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('CloudFront Distribution Report', 0)
            
            # Add metadata
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Target Account: {self.target_account_id}')
            doc.add_paragraph(f'Total Distributions: {len(distributions)}')
            doc.add_paragraph(f'Total Origin Access Identities: {len(oais)}')
            doc.add_paragraph('')
            
            # Define metric collection period (August 2025)
            start_time = datetime(2025, 8, 1, 0, 0, 0)
            end_time = datetime(2025, 8, 31, 23, 59, 59)
            
            # CloudFront Distributions with Metrics section
            doc.add_heading('CloudFront Distributions with Metrics', 1)
            if not distributions:
                doc.add_paragraph('No CloudFront distributions found.')
            else:
                for dist in distributions:
                    # Add distribution ID as heading
                    distribution_id = dist.get('Id')
                    doc.add_heading(f'{distribution_id}', 2)
                    
                    # Add basic distribution info
                    doc.add_paragraph(f"Domain Name: {dist.get('DomainName')}")

                    # Add Aliases if present
                    if dist.get('Aliases'):
                        doc.add_paragraph(f"Aliases: {', '.join(dist.get('Aliases'))}")
                        
                    doc.add_paragraph(f"Enabled: {self._safe_str(dist.get('Enabled'))}")  
                    
                    # # Add Origins info
                    # origins = dist.get('Origins', [])
                    # if origins:
                    #     doc.add_paragraph(f"Origins Count: {len(origins)}")
                    #     for origin in origins:
                    #         doc.add_paragraph(f"  - {origin.get('Id')}: {origin.get('DomainName')} ({origin.get('Type', 'Unknown')})", style='List Bullet')
                    
                    # doc.add_paragraph('')
                    
                    # Get and add Requests metric graph
                    doc.add_paragraph('Requests Metric :', style='Heading 3')
                    requests_data = self.get_cloudfront_metrics(distribution_id, 'Requests', start_time, end_time)
                    
                    # Add max value info if available
                    if requests_data['max_value'] is not None:
                        max_value_str = self._format_metric_value(requests_data['max_value'], 'Requests')
                        max_time_str = requests_data['max_timestamp'].strftime('%Y-%m-%d %H:%M')
                        doc.add_paragraph(f"최고값: {max_value_str} ({max_time_str})")
                    
                    if requests_data['datapoints'] or True:  # Always generate graph even if no data
                        requests_graph = self.generate_metric_graph(
                            distribution_id, 
                            'Requests', 
                            requests_data['datapoints'], 
                            requests_data['max_value'], 
                            requests_data['max_timestamp'], 
                            'Count'
                        )
                        if requests_graph.getbuffer().nbytes > 0:
                            doc.add_picture(requests_graph, width=Inches(6))
                        else:
                            doc.add_paragraph('Failed to generate Requests graph')
                    
                    doc.add_paragraph('')
                    
                    # Get and add BytesDownloaded metric graph
                    doc.add_paragraph('Bytes Downloaded Metric :', style='Heading 3')
                    bytes_data = self.get_cloudfront_metrics(distribution_id, 'BytesDownloaded', start_time, end_time)
                    
                    # Add max value info if available
                    if bytes_data['max_value'] is not None:
                        max_value_str = self._format_metric_value(bytes_data['max_value'], 'BytesDownloaded')
                        max_time_str = bytes_data['max_timestamp'].strftime('%Y-%m-%d %H:%M')
                        doc.add_paragraph(f"최고값: {max_value_str} ({max_time_str})")
                    
                    if bytes_data['datapoints'] or True:  # Always generate graph even if no data
                        bytes_graph = self.generate_metric_graph(
                            distribution_id, 
                            'BytesDownloaded', 
                            bytes_data['datapoints'], 
                            bytes_data['max_value'], 
                            bytes_data['max_timestamp'], 
                            'Bytes'
                        )
                        if bytes_graph.getbuffer().nbytes > 0:
                            doc.add_picture(bytes_graph, width=Inches(6))
                        else:
                            doc.add_paragraph('Failed to generate BytesDownloaded graph')
                    
                    # Add page break after each distribution (except the last one)
                    if dist != distributions[-1]:
                        doc.add_page_break()
            
            # Origin Access Identities section on new page
            if distributions:
                doc.add_page_break()
            
            doc.add_heading('Origin Access Identities', 1)
            if not oais:
                doc.add_paragraph('No Origin Access Identities found.')
            else:
                for oai in oais:
                    doc.add_paragraph(f"OAI ID: {self._safe_str(oai.get('Id'))}")
                    doc.add_paragraph(f"Comment: {self._safe_str(oai.get('Comment'))}")
                    doc.add_paragraph(f"S3 Canonical User ID: {self._safe_str(oai.get('S3CanonicalUserId'))[:50]}...")
                    doc.add_paragraph('')
            
            # Generate filename with current timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            filename = f"CloudFront_{timestamp}.docx"
            
            # Create temporary file
            temp_dir = tempfile.mkdtemp()
            file_path = os.path.join(temp_dir, filename)
            
            # Save document
            doc.save(file_path)
            logger.info(f"CloudFront Word report generated: {file_path}")
            
            return file_path
            
        except Exception as e:
            logger.error(f"Failed to generate CloudFront Word report: {str(e)}")
            raise
    
    def upload_to_s3(self, file_path: str) -> str:
        """
        Upload file to S3 bucket
        
        Args:
            file_path: Path to local file to upload
            
        Returns:
            S3 object key (path)
        """
        try:
            # Create S3 client using instance profile credentials
            s3_client = boto3.client('s3', region_name=self.region)
            
            # Extract filename from path
            filename = os.path.basename(file_path)
            
            # Construct S3 key
            s3_key = f"{settings.s3_report_folder}/{filename}"
            
            logger.info(f"Uploading file to s3://{settings.s3_bucket_name}/{s3_key}")
            
            # Upload file
            s3_client.upload_file(
                file_path,
                settings.s3_bucket_name,
                s3_key
            )
            
            logger.info(f"Successfully uploaded file to S3: {s3_key}")
            return s3_key
            
        except Exception as e:
            logger.error(f"Failed to upload file to S3: {str(e)}")
            raise
    
    def generate_and_upload_cloudfront_report(self) -> str:
        """
        Generate CloudFront Word report and upload to S3
        
        Returns:
            S3 object key of uploaded file
        """
        try:
            # Get CloudFront distributions
            distributions = self.list_cloudfront_distributions()
            
            # Get CloudFront OAIs
            oais = self.list_cloudfront_origin_access_identities()
            
            # Generate Word report
            file_path = self.generate_cloudfront_word_report(distributions, oais)
            
            # Upload to S3
            s3_key = self.upload_to_s3(file_path)
            
            # Clean up temporary file
            try:
                os.remove(file_path)
                os.rmdir(os.path.dirname(file_path))
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up temporary file: {cleanup_error}")
            
            return s3_key
            
        except Exception as e:
            logger.error(f"Failed to generate and upload CloudFront report: {str(e)}")
            raise


def main():
    """Main function to demonstrate CloudFront cross-account listing with metrics"""
    try:
        # Initialize the manager
        cf_manager = CloudFrontCrossAccountManager()
        
        # List all CloudFront distributions
        print(f"\n=== Listing all CloudFront distributions in account {settings.target_account_id} ===")
        distributions = cf_manager.list_cloudfront_distributions()
        
        # Display distributions in new format
        for dist in distributions:
            print(f"\n{dist['Id']}")
            print(f"  Domain Name: {dist['DomainName']}")
            print(f"  Status: {dist['Status']}")
            print(f"  Enabled: {dist['Enabled']}")
            if dist.get('Aliases'):
                print(f"  Aliases: {', '.join(dist['Aliases'])}")
            print(f"  Origins: {len(dist['Origins'])}")
            for origin in dist['Origins']:
                print(f"    - {origin.get('Id')}: {origin.get('DomainName')} ({origin.get('Type', 'Unknown')})")
            print(f"  Metrics will be included in the Word report")
        
        print(f"\nTotal CloudFront distributions: {len(distributions)}")
        
        # List all OAIs
        print(f"\n=== Listing all Origin Access Identities in account {settings.target_account_id} ===")
        oais = cf_manager.list_cloudfront_origin_access_identities()
        
        for oai in oais:
            print(f"\n{oai['Id']}")
            print(f"  Comment: {oai['Comment']}")
        
        print(f"\nTotal OAIs: {len(oais)}")
        
        # Generate and upload report with metrics
        print(f"\n=== Generating CloudFront Word report with CloudWatch metrics ===")
        print(f"Collecting metrics for period: 2025-08-01 to 2025-08-31")
        print(f"Metrics: Requests (Per-Distribution), BytesDownloaded (Per-Distribution)")
        print(f"Period: 1 Hour, Statistic: Sum")
        
        s3_key = cf_manager.generate_and_upload_cloudfront_report()
        print(f"\nCloudFront report with metrics uploaded to: s3://{settings.s3_bucket_name}/{s3_key}")
        
    except Exception as e:
        logger.error(f"Error in main: {str(e)}")
        raise


if __name__ == "__main__":
    main()