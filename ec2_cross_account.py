import boto3
import logging
from typing import List, Dict, Any
from botocore.exceptions import ClientError, BotoCoreError
from config import settings
from datetime import datetime
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Configure logging
logging.basicConfig(
    level=getattr(logging, settings.log_level),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class EC2CrossAccountManager:
    """Manages cross-account EC2 instance listing using AWS STS AssumeRole"""
    
    def __init__(self):
        self.source_account_id = settings.source_account_id
        self.target_account_id = settings.target_account_id
        self.assume_role_name = settings.assume_role_name
        self.session_name = settings.session_name
        self.region = settings.aws_region
        
    def assume_role(self) -> Dict[str, Any]:
        """
        Assume role in the target account
        
        Returns:
            Dict containing temporary credentials
        """
        try:
            # Create STS client using instance profile credentials
            sts_client = boto3.client('sts', region_name=self.region)
            
            # Construct the role ARN
            role_arn = f"arn:aws:iam::{self.target_account_id}:role/{self.assume_role_name}"
            
            logger.info(f"Attempting to assume role: {role_arn}")
            
            # Assume the role
            response = sts_client.assume_role(
                RoleArn=role_arn,
                RoleSessionName=self.session_name,
                DurationSeconds=3600,  # 1 hour
                ExternalId='starbucks-monitoring-secret-key-prod'  # Uncomment if using ExternalId
            )
            
            logger.info("Successfully assumed role in target account")
            return response['Credentials']
            
        except ClientError as e:
            error_code = e.response['Error']['Code']
            error_message = e.response['Error']['Message']
            logger.error(f"Failed to assume role: {error_code} - {error_message}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error during assume role: {str(e)}")
            raise
    
    def create_ec2_client(self, credentials: Dict[str, Any]):
        """
        Create EC2 client with assumed role credentials
        
        Args:
            credentials: Temporary credentials from assume_role
            
        Returns:
            EC2 client object
        """
        return boto3.client(
            'ec2',
            region_name=self.region,
            aws_access_key_id=credentials['AccessKeyId'],
            aws_secret_access_key=credentials['SecretAccessKey'],
            aws_session_token=credentials['SessionToken']
        )
    
    def create_elb_client(self, credentials: Dict[str, Any]):
        """
        Create ELB client with assumed role credentials
        
        Args:
            credentials: Temporary credentials from assume_role
            
        Returns:
            ELB client object
        """
        return boto3.client(
            'elbv2',
            region_name=self.region,
            aws_access_key_id=credentials['AccessKeyId'],
            aws_secret_access_key=credentials['SecretAccessKey'],
            aws_session_token=credentials['SessionToken']
        )
    
    def create_classic_elb_client(self, credentials: Dict[str, Any]):
        """
        Create Classic ELB client with assumed role credentials
        
        Args:
            credentials: Temporary credentials from assume_role
            
        Returns:
            Classic ELB client object
        """
        return boto3.client(
            'elb',
            region_name=self.region,
            aws_access_key_id=credentials['AccessKeyId'],
            aws_secret_access_key=credentials['SecretAccessKey'],
            aws_session_token=credentials['SessionToken']
        )
    
    def list_ec2_instances(self) -> List[Dict[str, Any]]:
        """
        List all EC2 instances in the target account
        
        Returns:
            List of EC2 instance information
        """
        try:
            # Assume role in target account
            credentials = self.assume_role()
            
            # Create EC2 client with assumed credentials
            ec2_client = self.create_ec2_client(credentials)
            
            # Describe all instances
            logger.info(f"Retrieving EC2 instances from account {self.target_account_id}")
            
            instances = []
            paginator = ec2_client.get_paginator('describe_instances')
            
            for page in paginator.paginate():
                for reservation in page['Reservations']:
                    for instance in reservation['Instances']:
                        instance_info = {
                            'InstanceId': instance['InstanceId'],
                            'InstanceType': instance['InstanceType'],
                            'State': instance['State']['Name'],
                            'LaunchTime': instance.get('LaunchTime', '').isoformat() if instance.get('LaunchTime') else None,
                            'PrivateIpAddress': instance.get('PrivateIpAddress'),
                            'PublicIpAddress': instance.get('PublicIpAddress'),
                            'Tags': self._parse_tags(instance.get('Tags', [])),
                            'VpcId': instance.get('VpcId'),
                            'SubnetId': instance.get('SubnetId'),
                            'AvailabilityZone': instance.get('Placement', {}).get('AvailabilityZone')
                        }
                        instances.append(instance_info)
            
            logger.info(f"Successfully retrieved {len(instances)} instances")
            return instances
            
        except Exception as e:
            logger.error(f"Failed to list EC2 instances: {str(e)}")
            raise
    
    def _parse_tags(self, tags: List[Dict[str, str]]) -> Dict[str, str]:
        """
        Parse EC2 tags into a dictionary
        
        Args:
            tags: List of tag dictionaries from EC2
            
        Returns:
            Dictionary of tag key-value pairs
        """
        return {tag['Key']: tag['Value'] for tag in tags}
    
    def get_instances_by_tag(self, tag_key: str, tag_value: str) -> List[Dict[str, Any]]:
        """
        Get EC2 instances filtered by a specific tag
        
        Args:
            tag_key: Tag key to filter by
            tag_value: Tag value to filter by
            
        Returns:
            List of EC2 instances matching the tag
        """
        all_instances = self.list_ec2_instances()
        filtered_instances = [
            instance for instance in all_instances
            if instance['Tags'].get(tag_key) == tag_value
        ]
        
        logger.info(f"Found {len(filtered_instances)} instances with tag {tag_key}={tag_value}")
        return filtered_instances
    
    def get_running_instances(self) -> List[Dict[str, Any]]:
        """
        Get only running EC2 instances
        
        Returns:
            List of running EC2 instances
        """
        all_instances = self.list_ec2_instances()
        running_instances = [
            instance for instance in all_instances
            if instance['State'] == 'running'
        ]
        
        logger.info(f"Found {len(running_instances)} running instances")
        return running_instances
    
    def list_load_balancers(self) -> List[Dict[str, Any]]:
        """
        List all Load Balancers (ALB/NLB and Classic ELB) in the target account
        
        Returns:
            List of Load Balancer information
        """
        try:
            # Assume role in target account
            credentials = self.assume_role()
            
            # Create ELB clients
            elbv2_client = self.create_elb_client(credentials)
            elb_client = self.create_classic_elb_client(credentials)
            
            logger.info(f"Retrieving Load Balancers from account {self.target_account_id}")
            
            load_balancers = []
            
            # Get ALB/NLB (ELBv2)
            try:
                paginator = elbv2_client.get_paginator('describe_load_balancers')
                for page in paginator.paginate():
                    for lb in page['LoadBalancers']:
                        # Get tags for the load balancer
                        tags_response = elbv2_client.describe_tags(
                            ResourceArns=[lb['LoadBalancerArn']]
                        )
                        tags = {}
                        if tags_response['TagDescriptions']:
                            tags = {tag['Key']: tag['Value'] for tag in tags_response['TagDescriptions'][0]['Tags']}
                        
                        lb_info = {
                            'Name': lb['LoadBalancerName'],
                            'Type': lb['Type'].upper(),  # application, network, gateway
                            'Scheme': lb['Scheme'],
                            'State': lb['State']['Code'],
                            'DNSName': lb['DNSName'],
                            'CanonicalHostedZoneId': lb['CanonicalHostedZoneId'],
                            'CreatedTime': lb['CreatedTime'].isoformat() if lb.get('CreatedTime') else None,
                            'VpcId': lb.get('VpcId'),
                            'AvailabilityZones': ', '.join([az['ZoneName'] for az in lb.get('AvailabilityZones', [])]),
                            'Tags': tags,
                            'LoadBalancerArn': lb['LoadBalancerArn']
                        }
                        load_balancers.append(lb_info)
                        
            except Exception as e:
                logger.warning(f"Failed to retrieve ALB/NLB: {str(e)}")
            
            # Get Classic ELB
            try:
                paginator = elb_client.get_paginator('describe_load_balancers')
                for page in paginator.paginate():
                    for lb in page['LoadBalancers']:
                        # Get tags for classic ELB
                        tags_response = elb_client.describe_tags(
                            LoadBalancerNames=[lb['LoadBalancerName']]
                        )
                        tags = {}
                        if tags_response['TagDescriptions']:
                            tags = {tag['Key']: tag['Value'] for tag in tags_response['TagDescriptions'][0]['Tags']}
                        
                        lb_info = {
                            'Name': lb['LoadBalancerName'],
                            'Type': 'CLASSIC',
                            'Scheme': lb['Scheme'],
                            'State': 'active',  # Classic ELBs don't have state
                            'DNSName': lb['DNSName'],
                            'CanonicalHostedZoneId': lb.get('CanonicalHostedZoneId'),
                            'CreatedTime': lb['CreatedTime'].isoformat() if lb.get('CreatedTime') else None,
                            'VpcId': lb.get('VPCId'),
                            'AvailabilityZones': ', '.join(lb.get('AvailabilityZones', [])),
                            'Tags': tags,
                            'LoadBalancerArn': f"arn:aws:elasticloadbalancing:{self.region}:{self.target_account_id}:loadbalancer/{lb['LoadBalancerName']}"
                        }
                        load_balancers.append(lb_info)
                        
            except Exception as e:
                logger.warning(f"Failed to retrieve Classic ELB: {str(e)}")
            
            logger.info(f"Successfully retrieved {len(load_balancers)} load balancers")
            return load_balancers
            
        except Exception as e:
            logger.error(f"Failed to list Load Balancers: {str(e)}")
            raise
    
    def _safe_str(self, value: Any) -> str:
        """
        Safely convert any value to string, handling None values
        
        Args:
            value: Any value that needs to be converted to string
            
        Returns:
            String representation of the value, or empty string if None
        """
        if value is None:
            return ''
        return str(value)
    
    def generate_word_report(self, instances: List[Dict[str, Any]]) -> str:
        """
        Generate Word document report with EC2 instances data
        
        Args:
            instances: List of EC2 instance information
            
        Returns:
            Path to generated Word document
        """
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('EC2 Instances Report', 0)
            
            # Add metadata
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Target Account: {self.target_account_id}')
            doc.add_paragraph(f'Total Instances: {len(instances)}')
            doc.add_paragraph('')
            
            if not instances:
                doc.add_paragraph('No EC2 instances found.')
            else:
                # Add table
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                
                # Header row
                headers = ['Instance ID', 'Type', 'State', 'Name', 'Private IP', 'Public IP', 'VPC ID', 'AZ']
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                
                # Data rows
                for instance in instances:
                    row_cells = table.add_row().cells
                    # Safely convert all values to strings, handling None values
                    row_cells[0].text = self._safe_str(instance.get('InstanceId'))
                    row_cells[1].text = self._safe_str(instance.get('InstanceType'))
                    row_cells[2].text = self._safe_str(instance.get('State'))
                    row_cells[3].text = self._safe_str(instance.get('Tags', {}).get('Name'))
                    row_cells[4].text = self._safe_str(instance.get('PrivateIpAddress'))
                    row_cells[5].text = self._safe_str(instance.get('PublicIpAddress'))
                    row_cells[6].text = self._safe_str(instance.get('VpcId'))
                    row_cells[7].text = self._safe_str(instance.get('AvailabilityZone'))
            
            # Generate filename with current timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            filename = f"{timestamp}.docx"
            
            # Create temporary file
            temp_dir = tempfile.mkdtemp()
            file_path = os.path.join(temp_dir, filename)
            
            # Save document
            doc.save(file_path)
            logger.info(f"Word report generated: {file_path}")
            
            return file_path
            
        except Exception as e:
            logger.error(f"Failed to generate Word report: {str(e)}")
            raise
    
    def generate_elb_word_report(self, load_balancers: List[Dict[str, Any]]) -> str:
        """
        Generate Word document report with ELB data
        
        Args:
            load_balancers: List of Load Balancer information
            
        Returns:
            Path to generated Word document
        """
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Load Balancers Report', 0)
            
            # Add metadata
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Target Account: {self.target_account_id}')
            doc.add_paragraph(f'Total Load Balancers: {len(load_balancers)}')
            doc.add_paragraph('')
            
            if not load_balancers:
                doc.add_paragraph('No Load Balancers found.')
            else:
                # Add table
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                
                # Header row
                headers = ['Name', 'Type', 'Scheme', 'State', 'DNS Name', 'VPC ID', 'AZs', 'Created']
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                
                # Data rows
                for lb in load_balancers:
                    row_cells = table.add_row().cells
                    # Safely convert all values to strings, handling None values
                    row_cells[0].text = self._safe_str(lb.get('Name'))
                    row_cells[1].text = self._safe_str(lb.get('Type'))
                    row_cells[2].text = self._safe_str(lb.get('Scheme'))
                    row_cells[3].text = self._safe_str(lb.get('State'))
                    row_cells[4].text = self._safe_str(lb.get('DNSName'))
                    row_cells[5].text = self._safe_str(lb.get('VpcId'))
                    row_cells[6].text = self._safe_str(lb.get('AvailabilityZones'))
                    # Format creation time to show only date
                    created_time = lb.get('CreatedTime')
                    if created_time:
                        try:
                            from datetime import datetime as dt
                            parsed_time = dt.fromisoformat(created_time.replace('Z', '+00:00'))
                            row_cells[7].text = parsed_time.strftime('%Y-%m-%d')
                        except:
                            row_cells[7].text = self._safe_str(created_time)[:10]
                    else:
                        row_cells[7].text = ''
            
            # Generate filename with current timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            filename = f"ELB_{timestamp}.docx"
            
            # Create temporary file
            temp_dir = tempfile.mkdtemp()
            file_path = os.path.join(temp_dir, filename)
            
            # Save document
            doc.save(file_path)
            logger.info(f"ELB Word report generated: {file_path}")
            
            return file_path
            
        except Exception as e:
            logger.error(f"Failed to generate ELB Word report: {str(e)}")
            raise
    
    def upload_to_s3(self, file_path: str) -> str:
        """
        Upload file to S3 bucket
        
        Args:
            file_path: Path to local file to upload
            
        Returns:
            S3 object key (path)
        """
        try:
            # Create S3 client using instance profile credentials
            s3_client = boto3.client('s3', region_name=self.region)
            
            # Extract filename from path
            filename = os.path.basename(file_path)
            
            # Construct S3 key
            s3_key = f"{settings.s3_report_folder}/{filename}"
            
            logger.info(f"Uploading file to s3://{settings.s3_bucket_name}/{s3_key}")
            
            # Upload file
            s3_client.upload_file(
                file_path,
                settings.s3_bucket_name,
                s3_key
            )
            
            logger.info(f"Successfully uploaded file to S3: {s3_key}")
            return s3_key
            
        except Exception as e:
            logger.error(f"Failed to upload file to S3: {str(e)}")
            raise
    
    def generate_and_upload_report(self) -> str:
        """
        Generate Word report and upload to S3
        
        Returns:
            S3 object key of uploaded file
        """
        try:
            # Get EC2 instances
            instances = self.list_ec2_instances()
            
            # Generate Word report
            file_path = self.generate_word_report(instances)
            
            # Upload to S3
            s3_key = self.upload_to_s3(file_path)
            
            # Clean up temporary file
            try:
                os.remove(file_path)
                os.rmdir(os.path.dirname(file_path))
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up temporary file: {cleanup_error}")
            
            return s3_key
            
        except Exception as e:
            logger.error(f"Failed to generate and upload report: {str(e)}")
            raise
    
    def generate_and_upload_elb_report(self) -> str:
        """
        Generate ELB Word report and upload to S3
        
        Returns:
            S3 object key of uploaded file
        """
        try:
            # Get Load Balancers
            load_balancers = self.list_load_balancers()
            
            # Generate Word report
            file_path = self.generate_elb_word_report(load_balancers)
            
            # Upload to S3
            s3_key = self.upload_to_s3(file_path)
            
            # Clean up temporary file
            try:
                os.remove(file_path)
                os.rmdir(os.path.dirname(file_path))
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up temporary file: {cleanup_error}")
            
            return s3_key
            
        except Exception as e:
            logger.error(f"Failed to generate and upload ELB report: {str(e)}")
            raise


    """Main function to demonstrate EC2 cross-account listing"""
    try:
        # Initialize the manager
        ec2_manager = EC2CrossAccountManager()
        
        # List all instances
        print(f"\n=== Listing all EC2 instances in account {settings.target_account_id} ===")
        instances = ec2_manager.list_ec2_instances()
        
        for instance in instances:
            print(f"\nInstance ID: {instance['InstanceId']}")
            print(f"  Type: {instance['InstanceType']}")
            print(f"  State: {instance['State']}")
            print(f"  Private IP: {instance['PrivateIpAddress']}")
            print(f"  Public IP: {instance['PublicIpAddress']}")
            print(f"  Name: {instance['Tags'].get('Name', 'N/A')}")
        
        # List only running instances
        print(f"\n=== Running instances only ===")
        running_instances = ec2_manager.get_running_instances()
        print(f"Total running instances: {len(running_instances)}")
        
        # Generate and upload EC2 report
        print(f"\n=== Generating and uploading EC2 Word report ===")
        ec2_s3_key = ec2_manager.generate_and_upload_report()
        print(f"EC2 Report uploaded to: s3://{settings.s3_bucket_name}/{ec2_s3_key}")
        
        # Generate and upload ELB report
        print(f"\n=== Generating and uploading ELB Word report ===")
        elb_s3_key = ec2_manager.generate_and_upload_elb_report()
        print(f"ELB Report uploaded to: s3://{settings.s3_bucket_name}/{elb_s3_key}")
        
        # List ELB information
        print(f"\n=== Load Balancers Summary ===")
        load_balancers = ec2_manager.list_load_balancers()
        print(f"Total Load Balancers: {len(load_balancers)}")
        
        # Show ELB details
        for lb in load_balancers[:5]:  # Show first 5 ELBs
            print(f"\nLoad Balancer: {lb.get('Name', 'N/A')}")
            print(f"  Type: {lb.get('Type', 'N/A')}")
            print(f"  State: {lb.get('State', 'N/A')}")
            print(f"  DNS: {lb.get('DNSName', 'N/A')}")
            print(f"  VPC: {lb.get('VpcId', 'N/A')}")
        
        if len(load_balancers) > 5:
            print(f"\n... and {len(load_balancers) - 5} more load balancers (see Word report for full details)")
        
    except Exception as e:
        logger.error(f"Error in main: {str(e)}")
        raise


if __name__ == "__main__":
    main()