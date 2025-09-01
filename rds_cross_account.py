import boto3
import logging
from typing import List, Dict, Any
from botocore.exceptions import ClientError, BotoCoreError
from config import settings
from datetime import datetime
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Configure logging
logging.basicConfig(
    level=getattr(logging, settings.log_level),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class RDSCrossAccountManager:
    """Manages cross-account RDS instance listing using AWS STS AssumeRole"""
    
    def __init__(self):
        self.source_account_id = settings.source_account_id
        self.target_account_id = settings.target_account_id
        self.assume_role_name = settings.assume_role_name
        self.session_name = settings.session_name
        self.region = settings.aws_region
        
    def assume_role(self) -> Dict[str, Any]:
        """
        Assume role in the target account
        
        Returns:
            Dict containing temporary credentials
        """
        try:
            # Create STS client using instance profile credentials
            sts_client = boto3.client('sts', region_name=self.region)
            
            # Construct the role ARN
            role_arn = f"arn:aws:iam::{self.target_account_id}:role/{self.assume_role_name}"
            
            logger.info(f"Attempting to assume role: {role_arn}")
            
            # Assume the role
            response = sts_client.assume_role(
                RoleArn=role_arn,
                RoleSessionName=self.session_name,
                DurationSeconds=3600,  # 1 hour
                # ExternalId='starbucks-monitoring-secret-key-prod'  # Uncomment if using ExternalId
            )
            
            logger.info("Successfully assumed role in target account")
            return response['Credentials']
            
        except ClientError as e:
            error_code = e.response['Error']['Code']
            error_message = e.response['Error']['Message']
            logger.error(f"Failed to assume role: {error_code} - {error_message}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error during assume role: {str(e)}")
            raise
    
    def create_rds_client(self, credentials: Dict[str, Any]):
        """
        Create RDS client with assumed role credentials
        
        Args:
            credentials: Temporary credentials from assume_role
            
        Returns:
            RDS client object
        """
        return boto3.client(
            'rds',
            region_name=self.region,
            aws_access_key_id=credentials['AccessKeyId'],
            aws_secret_access_key=credentials['SecretAccessKey'],
            aws_session_token=credentials['SessionToken']
        )
    
    def list_rds_instances(self) -> List[Dict[str, Any]]:
        """
        List all RDS instances in the target account
        
        Returns:
            List of RDS instance information
        """
        try:
            # Assume role in target account
            credentials = self.assume_role()
            
            # Create RDS client with assumed credentials
            rds_client = self.create_rds_client(credentials)
            
            # Describe all DB instances
            logger.info(f"Retrieving RDS instances from account {self.target_account_id}")
            
            instances = []
            paginator = rds_client.get_paginator('describe_db_instances')
            
            for page in paginator.paginate():
                for db_instance in page['DBInstances']:
                    instance_info = {
                        'DBInstanceIdentifier': db_instance.get('DBInstanceIdentifier'),
                        'DBInstanceClass': db_instance.get('DBInstanceClass'),
                        'Engine': db_instance.get('Engine'),
                        'EngineVersion': db_instance.get('EngineVersion'),
                        'DBInstanceStatus': db_instance.get('DBInstanceStatus'),
                        'Endpoint': db_instance.get('Endpoint', {}).get('Address'),
                        'Port': db_instance.get('Endpoint', {}).get('Port'),
                        'AllocatedStorage': db_instance.get('AllocatedStorage'),
                        'StorageType': db_instance.get('StorageType'),
                        'MultiAZ': db_instance.get('MultiAZ'),
                        'VpcId': db_instance.get('DBSubnetGroup', {}).get('VpcId'),
                        'AvailabilityZone': db_instance.get('AvailabilityZone'),
                        'BackupRetentionPeriod': db_instance.get('BackupRetentionPeriod'),
                        'PreferredBackupWindow': db_instance.get('PreferredBackupWindow'),
                        'PreferredMaintenanceWindow': db_instance.get('PreferredMaintenanceWindow'),
                        'Tags': self._parse_tags(db_instance.get('TagList', []))
                    }
                    instances.append(instance_info)
            
            logger.info(f"Successfully retrieved {len(instances)} RDS instances")
            return instances
            
        except Exception as e:
            logger.error(f"Failed to list RDS instances: {str(e)}")
            raise
    
    def list_rds_clusters(self) -> List[Dict[str, Any]]:
        """
        List all RDS clusters (Aurora) in the target account
        
        Returns:
            List of RDS cluster information
        """
        try:
            # Assume role in target account
            credentials = self.assume_role()
            
            # Create RDS client with assumed credentials
            rds_client = self.create_rds_client(credentials)
            
            # Describe all DB clusters
            logger.info(f"Retrieving RDS clusters from account {self.target_account_id}")
            
            clusters = []
            paginator = rds_client.get_paginator('describe_db_clusters')
            
            for page in paginator.paginate():
                for db_cluster in page['DBClusters']:
                    cluster_info = {
                        'DBClusterIdentifier': db_cluster.get('DBClusterIdentifier'),
                        'Engine': db_cluster.get('Engine'),
                        'EngineVersion': db_cluster.get('EngineVersion'),
                        'Status': db_cluster.get('Status'),
                        'Endpoint': db_cluster.get('Endpoint'),
                        'ReaderEndpoint': db_cluster.get('ReaderEndpoint'),
                        'Port': db_cluster.get('Port'),
                        'MultiAZ': db_cluster.get('MultiAZ'),
                        'AllocatedStorage': db_cluster.get('AllocatedStorage'),
                        'BackupRetentionPeriod': db_cluster.get('BackupRetentionPeriod'),
                        'PreferredBackupWindow': db_cluster.get('PreferredBackupWindow'),
                        'PreferredMaintenanceWindow': db_cluster.get('PreferredMaintenanceWindow'),
                        'VpcId': db_cluster.get('DBSubnetGroup', {}).get('VpcId') if db_cluster.get('DBSubnetGroup') else None,
                        'Tags': self._parse_tags(db_cluster.get('TagList', []))
                    }
                    clusters.append(cluster_info)
            
            logger.info(f"Successfully retrieved {len(clusters)} RDS clusters")
            return clusters
            
        except Exception as e:
            logger.error(f"Failed to list RDS clusters: {str(e)}")
            raise
    
    def _parse_tags(self, tags: List[Dict[str, str]]) -> Dict[str, str]:
        """
        Parse RDS tags into a dictionary
        
        Args:
            tags: List of tag dictionaries from RDS
            
        Returns:
            Dictionary of tag key-value pairs
        """
        return {tag.get('Key', ''): tag.get('Value', '') for tag in tags}
    
    def _safe_str(self, value: Any) -> str:
        """
        Safely convert any value to string, handling None values
        
        Args:
            value: Any value that needs to be converted to string
            
        Returns:
            String representation of the value, or empty string if None
        """
        if value is None:
            return ''
        if isinstance(value, bool):
            return 'Yes' if value else 'No'
        return str(value)
    
    def generate_rds_word_report(self, instances: List[Dict[str, Any]], clusters: List[Dict[str, Any]]) -> str:
        """
        Generate Word document report with RDS instances and clusters data
        
        Args:
            instances: List of RDS instance information
            clusters: List of RDS cluster information
            
        Returns:
            Path to generated Word document
        """
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('RDS Database Report', 0)
            
            # Add metadata
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Target Account: {self.target_account_id}')
            doc.add_paragraph(f'Total RDS Instances: {len(instances)}')
            doc.add_paragraph(f'Total RDS Clusters: {len(clusters)}')
            doc.add_paragraph('')
            
            # RDS Instances section
            doc.add_heading('RDS Instances', 1)
            if not instances:
                doc.add_paragraph('No RDS instances found.')
            else:
                # Add table for instances
                table = doc.add_table(rows=1, cols=9)
                table.style = 'Table Grid'
                
                # Header row
                headers = ['DB Instance', 'Class', 'Engine', 'Status', 'Endpoint', 'Storage (GB)', 'Multi-AZ', 'Backup Days', 'Name Tag']
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                
                # Data rows
                for instance in instances:
                    row_cells = table.add_row().cells
                    row_cells[0].text = self._safe_str(instance.get('DBInstanceIdentifier'))
                    row_cells[1].text = self._safe_str(instance.get('DBInstanceClass'))
                    row_cells[2].text = f"{self._safe_str(instance.get('Engine'))} {self._safe_str(instance.get('EngineVersion'))}"
                    row_cells[3].text = self._safe_str(instance.get('DBInstanceStatus'))
                    row_cells[4].text = self._safe_str(instance.get('Endpoint'))
                    row_cells[5].text = self._safe_str(instance.get('AllocatedStorage'))
                    row_cells[6].text = self._safe_str(instance.get('MultiAZ'))
                    row_cells[7].text = self._safe_str(instance.get('BackupRetentionPeriod'))
                    row_cells[8].text = self._safe_str(instance.get('Tags', {}).get('Name'))
            
            doc.add_paragraph('')
            
            # RDS Clusters section
            doc.add_heading('RDS Clusters (Aurora)', 1)
            if not clusters:
                doc.add_paragraph('No RDS clusters found.')
            else:
                # Add table for clusters
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                
                # Header row
                headers = ['Cluster ID', 'Engine', 'Status', 'Writer Endpoint', 'Reader Endpoint', 'Multi-AZ', 'Backup Days', 'Name Tag']
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                
                # Data rows
                for cluster in clusters:
                    row_cells = table.add_row().cells
                    row_cells[0].text = self._safe_str(cluster.get('DBClusterIdentifier'))
                    row_cells[1].text = f"{self._safe_str(cluster.get('Engine'))} {self._safe_str(cluster.get('EngineVersion'))}"
                    row_cells[2].text = self._safe_str(cluster.get('Status'))
                    row_cells[3].text = self._safe_str(cluster.get('Endpoint'))
                    row_cells[4].text = self._safe_str(cluster.get('ReaderEndpoint'))
                    row_cells[5].text = self._safe_str(cluster.get('MultiAZ'))
                    row_cells[6].text = self._safe_str(cluster.get('BackupRetentionPeriod'))
                    row_cells[7].text = self._safe_str(cluster.get('Tags', {}).get('Name'))
            
            # Generate filename with current timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            filename = f"RDS_{timestamp}.docx"
            
            # Create temporary file
            temp_dir = tempfile.mkdtemp()
            file_path = os.path.join(temp_dir, filename)
            
            # Save document
            doc.save(file_path)
            logger.info(f"RDS Word report generated: {file_path}")
            
            return file_path
            
        except Exception as e:
            logger.error(f"Failed to generate RDS Word report: {str(e)}")
            raise
    
    def upload_to_s3(self, file_path: str) -> str:
        """
        Upload file to S3 bucket
        
        Args:
            file_path: Path to local file to upload
            
        Returns:
            S3 object key (path)
        """
        try:
            # Create S3 client using instance profile credentials
            s3_client = boto3.client('s3', region_name=self.region)
            
            # Extract filename from path
            filename = os.path.basename(file_path)
            
            # Construct S3 key
            s3_key = f"{settings.s3_report_folder}/{filename}"
            
            logger.info(f"Uploading file to s3://{settings.s3_bucket_name}/{s3_key}")
            
            # Upload file
            s3_client.upload_file(
                file_path,
                settings.s3_bucket_name,
                s3_key
            )
            
            logger.info(f"Successfully uploaded file to S3: {s3_key}")
            return s3_key
            
        except Exception as e:
            logger.error(f"Failed to upload file to S3: {str(e)}")
            raise
    
    def generate_and_upload_rds_report(self) -> str:
        """
        Generate RDS Word report and upload to S3
        
        Returns:
            S3 object key of uploaded file
        """
        try:
            # Get RDS instances
            instances = self.list_rds_instances()
            
            # Get RDS clusters
            clusters = self.list_rds_clusters()
            
            # Generate Word report
            file_path = self.generate_rds_word_report(instances, clusters)
            
            # Upload to S3
            s3_key = self.upload_to_s3(file_path)
            
            # Clean up temporary file
            try:
                os.remove(file_path)
                os.rmdir(os.path.dirname(file_path))
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up temporary file: {cleanup_error}")
            
            return s3_key
            
        except Exception as e:
            logger.error(f"Failed to generate and upload RDS report: {str(e)}")
            raise


def main():
    """Main function to demonstrate RDS cross-account listing"""
    try:
        # Initialize the manager
        rds_manager = RDSCrossAccountManager()
        
        # List all RDS instances
        print(f"\n=== Listing all RDS instances in account {settings.target_account_id} ===")
        instances = rds_manager.list_rds_instances()
        
        for instance in instances:
            print(f"\nDB Instance: {instance['DBInstanceIdentifier']}")
            print(f"  Class: {instance['DBInstanceClass']}")
            print(f"  Engine: {instance['Engine']} {instance['EngineVersion']}")
            print(f"  Status: {instance['DBInstanceStatus']}")
            print(f"  Endpoint: {instance['Endpoint']}")
            print(f"  Multi-AZ: {instance['MultiAZ']}")
            print(f"  Name: {instance['Tags'].get('Name', 'N/A')}")
        
        print(f"\nTotal RDS instances: {len(instances)}")
        
        # List all RDS clusters
        print(f"\n=== Listing all RDS clusters in account {settings.target_account_id} ===")
        clusters = rds_manager.list_rds_clusters()
        
        for cluster in clusters:
            print(f"\nDB Cluster: {cluster['DBClusterIdentifier']}")
            print(f"  Engine: {cluster['Engine']} {cluster['EngineVersion']}")
            print(f"  Status: {cluster['Status']}")
            print(f"  Endpoint: {cluster['Endpoint']}")
            print(f"  Name: {cluster['Tags'].get('Name', 'N/A')}")
        
        print(f"\nTotal RDS clusters: {len(clusters)}")
        
        # Generate and upload report
        print(f"\n=== Generating and uploading RDS Word report ===")
        s3_key = rds_manager.generate_and_upload_rds_report()
        print(f"RDS report uploaded to: s3://{settings.s3_bucket_name}/{s3_key}")
        
    except Exception as e:
        logger.error(f"Error in main: {str(e)}")
        raise


if __name__ == "__main__":
    main()